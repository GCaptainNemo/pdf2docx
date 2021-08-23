from PIL import Image
# import pytesseract
import os
import re
import shutil
from docx import Document
import cv2


# document = Document(docx="NewYorkCourseMateirals.docx")
document = Document()
img_files = os.listdir("image/")
for i in range(195):
    if not 3 <= i <= 14:
        continue
    path = "image/{}.png".format(i)
    if os.path.exists(path):
        img = Image.open(path)
    text = pytesseract.image_to_string(img, lang='eng')
    print(text.count("\n"))
    text = text.replace("\r", "\n")
    text_lst = text.strip().split("\n")
    while "" in text_lst:
        text_lst.remove("")
    print(text_lst)
    for tile in text_lst:
        if not tile or tile.isspace():
            continue
        document.add_paragraph(tile.encode('utf-8').decode("utf-8"))
    document.add_heading("Page {}/195".format(i + 1))
    document.add_page_break()
    # with open("NewYorkCourseMateirals.txt", "w+", encoding='utf-8') as f:
    #     f.write(text)
    # print("########################################################3")
    # print(text)
    # print(type(text))

document.save("Content.docx")
#
