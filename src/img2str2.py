import easyocr
import cv2
from docx import Document
import pickle
import os
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

PROPORTION = 1
reader = easyocr.Reader(['en'], gpu=True)


class Dealer:
    @staticmethod
    def deal_result(result):
        """
        deal with result according to detection position
        :return:
        row_content_lst: 每行合并的内容
        delta_y_lst: 行与行之间距离
        delta_x_lst: 每行距离最右端的距离

        """
        result.sort(key=lambda x: x[0][0][1])  # sort according to y
        min_x_pos = min(_[0][0][0] for _ in result)
        max_x_pos = max(_[0][0][0] for _ in result)
        min_y_pos = min(_[0][0][1] for _ in result)
        print("min_x_pos = ", min_x_pos)
        print("min_y_pos = ", min_y_pos)

        rows_lst = []
        for idx in range(len(result)):
            if idx == 0:
                row_lst = [result[idx]]
            elif abs((result[idx][0][0][1] + result[idx][0][3][1]) / 2 -
                     (row_lst[-1][0][0][1] + row_lst[-1][0][3][1]) / 2) < 4:
                row_lst.append(result[idx])
            else:
                row_lst.sort(key=lambda x: x[0][0][0])   # sort according to x
                rows_lst.append(row_lst)
                row_lst = [result[idx]]
        if row_lst:
            rows_lst.append(row_lst)
        for idx, row in enumerate(rows_lst):
            print("row {} = ".format(idx), row)
        # #################################################
        # delta_x_lst: 每行距离最右端的距离
        # row_content_lst: 每行合并的内容
        # delta_y_lst: 行与行之间距离
        # ###################################################
        row_content_lst = []
        delta_x_lst = []
        for row_idx, row in enumerate(rows_lst):
            delta_x_lst.append(max_x_pos - row[-1][0][1][0])
            # ##############################################################3
            row_content = ""
            for idx, content in enumerate(row):
                if idx == 0:
                    row_content += " " * int((content[0][0][0] - min_x_pos) // 30 + 1) + content[1]
                else:
                    row_content += " " * int((row[idx][0][0][0] - row[idx - 1][0][1][0]) // 30 + 1) + \
                                  content[1]
            row_content_lst.append(row_content)
            print("row {} = ".format(row_idx), row_content)

        delta_y_lst = []
        for row_idx in range(len(rows_lst) - 1):
            delta_y = rows_lst[row_idx + 1][0][0][0][1] - rows_lst[row_idx][0][0][3][1]
            delta_y_lst.append(delta_y)

        print("delta_y_lst = ", delta_y_lst)
        return row_content_lst, delta_y_lst, delta_x_lst

    @staticmethod
    def row_deltay2par(row_content_lst, delta_y_lst, delta_x_lst):
        par_lst = []
        par_str = ""
        for i in range(len(row_content_lst)):
            if i == 0:
                par_str += row_content_lst[i]
            elif delta_y_lst[i - 1] < 14 and delta_x_lst[i] < 10:
                # 行之间距离小于14，距离末尾距离小于10，则认为是一段
                par_str += row_content_lst[i]
            else:
                # par_lst.append(int(delta_y_lst[i - 1] / 10) * "\n" + par_str)
                # par_lst.append(par_str + int(delta_y_lst[i - 1] / 10) * "\n")
                par_lst.append(par_str)

                par_str = row_content_lst[i]
        par_lst.append(par_str)
        return par_lst


class IoOBJ:
    @staticmethod
    def img2docx(par_lst, out_address):
        document = Document()
        document.styles['Normal'].font.name = u'Times New Roman'
        document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'Times New Roman')
        document.styles['Normal'].font.size = Pt(10.5)  # 5号字
        document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

        for row_content in par_lst:
            document.add_paragraph(row_content.encode('utf-8').decode("utf-8"))
        document.save(out_address)

    @staticmethod
    def imgs_to_docx(proportion):
        document = Document()
        document.styles['Normal'].font.name = u'Times New Roman'
        document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'Times New Roman')
        document.styles['Normal'].font.size = Pt(10.5)  # 5号字
        document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        # Head = document.add_heading("", level=1)  # 这里不填标题内容
        # run = Head.add_run("刚来csdn，这就是博客么，I了")
        # run.font.name = u'Cambria'
        # run.font.color.rgb = RGBColor(0, 0, 0)
        # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
        for i in range(195):
            with open("../result/{}_{}.pkl".format(i, proportion), "rb+") as f:
                result = pickle.load(f)
            row_content_lst, delta_y_lst, delta_x_lst = Dealer.deal_result(result)
            par_lst = Dealer.row_deltay2par(row_content_lst, delta_y_lst, delta_x_lst)
            for row_content in par_lst:
                document.add_paragraph(row_content.encode('utf-8').decode("utf-8"))
            document.add_heading("Page {}/195".format(i + 1))
            document.add_page_break()
        document.save("new_edition.docx")


class Utils:
    @staticmethod
    def dump_imgs(input_address, output_address, proportion):
        files_dir = os.listdir(input_address)
        if not os.path.exists(output_address):
            os.makedirs(output_address)
        for i, file in enumerate(files_dir):
            img = cv2.imread('../image/{}'.format(file))
            img = cv2.resize(img, (int(img.shape[1] / proportion), int(img.shape[0] / proportion)),
                             cv2.INTER_CUBIC)
            dump_file_name = "{}_{}.png".format(file.split(".")[0], proportion)
            cv2.imwrite(output_address + "/" + dump_file_name, img=img)
            print("dump image" + dump_file_name)

    @staticmethod
    def dump_result(in_address, out_address):
        files_dir = os.listdir(in_address)
        if not os.path.exists(out_address):
            os.makedirs(out_address)
        for i, file in enumerate(files_dir):
            result = reader.readtext(in_address + "/" + file)
            dump_res_name = file.split(".")[0] + ".pkl"
            with open(out_address + "/" + dump_res_name, "wb+") as f:
                pickle.dump(result, f)
            print("dump result {}".format(dump_res_name))

    @staticmethod
    def show_detect_img(result, img_address):
        img = cv2.imread(img_address)
        for i, (coordinate, txt, prob) in enumerate(result):
            for index in range(4):
                if index < 3:
                    img = cv2.line(img, (int(coordinate[index][0]), int(coordinate[index][1])),
                                   (int(coordinate[index + 1][0]), int(coordinate[index + 1][1])), color=(255, 0, 0))
                else:
                    img = cv2.line(img, (int(coordinate[index][0]), int(coordinate[index][1])),
                                   (int(coordinate[0][0]), int(coordinate[0][1])), color=(255, 0, 0))
        cv2.namedWindow("img", cv2.WINDOW_NORMAL)
        cv2.imshow("img", img)
        cv2.waitKey(0)


if __name__ == "__main__":
    # img = cv2.imread('../image/24.png')
    # img = cv2.resize(img, (int(img.shape[1] / PROPORTION), int(img.shape[0] / PROPORTION)), cv2.INTER_CUBIC)
    # cv2.imwrite("../image/24_.png", img=img)
    # res = reader.readtext('../image/24.png')
    # row_content_lst, delta_y_lst, delta_x_lst = Dealer.deal_result(result=res)
    # par_lst = Dealer.row_deltay2par(row_content_lst, delta_y_lst)
    # IoOBJ.img2docx(par_lst, "../test.docx")
    # Utils.show_detect_img(result=res, img_address="../image/24.png")
    # Utils.dump_imgs("../image/", output_address="../linshi/p4", proportion=4)
    Utils.dump_result("../linshi/p4", "../result/p4")
